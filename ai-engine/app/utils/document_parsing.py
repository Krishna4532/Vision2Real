"""Utilities for parsing documents (PDF, DOCX, TXT) and extracting text."""
from __future__ import annotations

import io
from pathlib import Path


async def extract_text_from_pdf(file_path: str | bytes) -> str:
    """Extract text from a PDF file.
    
    Args:
        file_path: Path to PDF file or bytes content
        
    Returns:
        Extracted text from PDF
    """
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
    
    try:
        if isinstance(file_path, bytes):
            pdf_file = io.BytesIO(file_path)
        else:
            pdf_file = open(file_path, 'rb')
        
        reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        if isinstance(file_path, str):
            pdf_file.close()
            
        return text.strip()
    except Exception as exc:
        raise ValueError(f"Failed to extract text from PDF: {exc}") from exc


async def extract_text_from_docx(file_path: str | bytes) -> str:
    """Extract text from a DOCX file.
    
    Args:
        file_path: Path to DOCX file or bytes content
        
    Returns:
        Extracted text from DOCX
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    
    try:
        if isinstance(file_path, bytes):
            doc = Document(io.BytesIO(file_path))
        else:
            doc = Document(file_path)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as exc:
        raise ValueError(f"Failed to extract text from DOCX: {exc}") from exc


async def extract_text_from_file(file_path: str | bytes, file_type: str | None = None) -> str:
    """Extract text from a file based on its type.
    
    Args:
        file_path: Path to file or bytes content
        file_type: File type (pdf, docx, txt). If None, inferred from extension
        
    Returns:
        Extracted text
    """
    if file_type is None:
        if isinstance(file_path, bytes):
            raise ValueError("file_type must be specified when passing bytes content")
        file_type = Path(file_path).suffix.lower().lstrip('.')
    
    file_type = file_type.lower()
    
    if file_type == "pdf":
        return await extract_text_from_pdf(file_path)
    elif file_type in ("docx", "doc"):
        return await extract_text_from_docx(file_path)
    elif file_type == "txt":
        if isinstance(file_path, bytes):
            return file_path.decode("utf-8")
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {file_type}. Supported: pdf, docx, txt")


def get_document_metadata(content: str, file_type: str) -> dict[str, str]:
    """Extract basic metadata from document content.
    
    Args:
        content: Document text content
        file_type: Document file type
        
    Returns:
        Dictionary with metadata
    """
    lines = content.split('\n')
    
    # Try to find title (first meaningful line or first line with capitals)
    title = ""
    for line in lines[:10]:
        line = line.strip()
        if line and len(line) > 5 and len(line) < 200:
            title = line
            break
    
    return {
        "file_type": file_type,
        "title": title,
        "word_count": str(len(content.split())),
        "line_count": str(len(lines)),
    }
